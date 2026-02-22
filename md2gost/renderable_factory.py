import logging
from functools import singledispatchmethod

from docx.shared import Parented, RGBColor

from .renderable import *
from .renderable import Renderable
from .renderable.content_sequence import ContentSequence
from .renderable.image import Image as ImageRenderable
from . import extended_markdown
from .renderable.caption import CaptionInfo
from .renderable.paragraph import Link
from .renderable.table import Table
from .renderable.equation import Equation
from .renderable.heading import Heading
from .renderable.list import List
from .renderable.toc import ToC
from .renderable.page_break import PageBreak


class RenderableFactory:
    def __init__(self, parent: Parented):
        self._parent = parent

    @singledispatchmethod
    def create(self, marko_element: extended_markdown.BlockElement,
               caption_info: CaptionInfo) -> Renderable:
        paragraph = Paragraph(self._parent)
        paragraph.add_run(f"{marko_element.get_type()} is not supported", color=RGBColor.from_string('ff0000'))
        logging.warning(f"{marko_element.get_type()} is not supported")
        return paragraph

    @staticmethod
    def _add_child_to_paragraph(paragraph_or_link: Paragraph | Link, child, classes: list[type]):
        if isinstance(child, (extended_markdown.RawText, extended_markdown.Literal)):
            paragraph_or_link.add_run(child.children,
                                      is_bold=extended_markdown.StrongEmphasis in classes or None,
                                      is_italic=extended_markdown.Emphasis in classes or None,
                                      strike_through=extended_markdown.Strikethrough in classes or None)
        elif isinstance(child, extended_markdown.CodeSpan):
            paragraph_or_link.add_run(child.children, is_italic=True)
        elif isinstance(child, extended_markdown.Image):
            paragraph_or_link.add_image(child.dest, CaptionInfo(None, child.title))
        elif isinstance(child, extended_markdown.LineBreak):
            pass  # ignore
        elif isinstance(child, extended_markdown.InlineEquation):
            paragraph_or_link.add_inline_equation(child.latex_equation)
        elif isinstance(child, (extended_markdown.Link, extended_markdown.Url)):
            RenderableFactory._create_runs(paragraph_or_link.add_link(child.dest),
                                           child.children, classes)
        elif isinstance(child, (extended_markdown.Emphasis, extended_markdown.StrongEmphasis,
                                extended_markdown.Strikethrough)):
            RenderableFactory._create_runs(paragraph_or_link,
                                           child.children, classes + [type(child)])
        else:
            paragraph_or_link.add_run(f" {child.get_type()} is not supported ",
                                      color=RGBColor.from_string("FF0000"))
            logging.warning(f"{child.get_type()} is not supported")

    @staticmethod
    def _create_runs(paragraph_or_link: Paragraph | Link, children, classes: list[type] = None):
        if not classes:
            classes = []
        for child in children:
            RenderableFactory._add_child_to_paragraph(paragraph_or_link, child, classes)

    @staticmethod
    def _build_paragraph_content_items(parent: Parented, children) -> list[Renderable]:
        """Build list of Paragraph and Image renderables in source order.
        Ensures images appear between the correct text segments."""
        items: list[Renderable] = []
        current = Paragraph(parent)
        for child in children:
            if isinstance(child, extended_markdown.Image):
                if current._docx_paragraph.text.strip():
                    items.append(current)
                items.append(ImageRenderable(parent, child.dest, CaptionInfo(None, child.title)))
                current = Paragraph(parent)
            else:
                RenderableFactory._add_child_to_paragraph(current, child, [])
        if current._docx_paragraph.text.strip() or current._images:
            items.append(current)
        return items

    @create.register
    def _(self, marko_paragraph: extended_markdown.Paragraph, caption_info: CaptionInfo):
        items = RenderableFactory._build_paragraph_content_items(
            self._parent, marko_paragraph.children
        )
        if len(items) == 0:
            return Paragraph(self._parent)
        if len(items) == 1:
            return items[0]
        return ContentSequence(self._parent, items)

    @create.register
    def _(self, marko_heading: extended_markdown.Heading, caption_info: CaptionInfo):
        heading = Heading(self._parent, marko_heading.level, marko_heading.numbered)
        # Для нумерованных заголовков добавляем ведущий пробел, чтобы при двузначной/трёхзначной
        # нумерации (1.2.10, 1.2.100) между номером и текстом всегда был отступ.
        if marko_heading.numbered:
            heading.add_run(" ")
        RenderableFactory._create_runs(heading, marko_heading.children)
        return heading

    @create.register
    def _(self, marko_code_block: extended_markdown.FencedCode | extended_markdown.CodeBlock, caption_info: CaptionInfo):
        listing = Listing(self._parent, marko_code_block.lang, caption_info)
        listing.set_text(marko_code_block.children[0].children)
        return listing

    @create.register
    def _(self, marko_equation: extended_markdown.Equation, caption_info: CaptionInfo):
        formula = Equation(self._parent, marko_equation.latex_equation)
        return formula

    @create.register
    def _(self, marko_list: extended_markdown.List, caption_info: CaptionInfo):
        list_ = List(self._parent, marko_list.ordered)

        def create_items_from_marko(marko_list_, level=1):
            for list_item in marko_list_.children:
                for child in list_item.children:
                    if isinstance(child, extended_markdown.List):
                        create_items_from_marko(child, level + 1)
                    elif isinstance(child, extended_markdown.Paragraph):
                        RenderableFactory._create_runs(
                            list_.add_item(level),
                            child.children
                        )

        create_items_from_marko(marko_list)

        return list_

    @create.register
    def _(self, marko_table: extended_markdown.Table, caption_info: CaptionInfo):
        table = Table(self._parent, len(marko_table.children), len(marko_table.children[0].children),
                      caption_info)

        for i, row in enumerate(marko_table.children):
            for j, cell in enumerate(row.children):
                RenderableFactory._create_runs(
                    table.add_paragraph_to_cell(i, j),
                    cell.children
                )

        return table

    @create.register
    def _(self, marko_toc: extended_markdown.TOC, caption_info: CaptionInfo):
        toc = ToC(self._parent)
        return toc

    @create.register
    def _(self, marko_thematic_break: extended_markdown.ThematicBreak, caption_info: CaptionInfo):
        return PageBreak(self._parent)
