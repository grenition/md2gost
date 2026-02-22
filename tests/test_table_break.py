# -*- coding: utf-8 -*-
import unittest

try:
    import matplotlib  # noqa: F401 — нужен для расчёта высот в Caption/Paragraph
    _has_matplotlib = True
except ImportError:
    _has_matplotlib = False

from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.shared import Pt

from md2gost.layout_tracker import LayoutState
from md2gost.renderable.table import Table
from md2gost.renderable.caption import CaptionInfo
from md2gost.rendered_info import RenderedInfo

from . import _create_test_document


@unittest.skipUnless(_has_matplotlib, "тест требует matplotlib (расчёт высот шрифтов)")
class TestTableBreak(unittest.TestCase):
    """Перенос таблицы: разрыв должен происходить до строки, которая не помещается
    (table_height + row_height > remaining), а не только по высоте строки.
    """

    def setUp(self) -> None:
        self._document, self._max_height, self._max_width = _create_test_document()

    def test_table_break_uses_table_plus_row_height(self):
        """При нехватке места перенос делается до строки; первая часть таблицы
        содержит все строки, помещающиеся на странице; «Продолжение таблицы» — один раз.
        """
        parent = self._document._body
        caption_info = CaptionInfo(None, "Test")
        table = Table(parent, 2, 2, caption_info)
        table.add_paragraph_to_cell(0, 0).add_run("A")
        table.add_paragraph_to_cell(0, 1).add_run("B")
        table.add_paragraph_to_cell(1, 0).add_run("C")
        table.add_paragraph_to_cell(1, 1).add_run("D")
        table.set_number(1)

        # Оставляем мало места после заголовка: первая строка поместится,
        # вторая — нет (table_height + row_height > remaining).
        layout_state = LayoutState(self._max_height, self._max_width)
        layout_state.add_height(self._max_height - Pt(45))

        infos = [info for info in table.render(None, layout_state) if isinstance(info, RenderedInfo)]

        continuation = [
            i for i in infos
            if isinstance(i.docx_element, DocxParagraph) and "Продолжение таблицы" in (i.docx_element.text or "")
        ]
        self.assertEqual(len(continuation), 1, "Ожидалась ровно одна подпись «Продолжение таблицы»")

        table_parts = [i for i in infos if isinstance(i.docx_element, DocxTable)]
        self.assertGreaterEqual(len(table_parts), 2, "Ожидались минимум две части таблицы (до и после переноса)")
        self.assertEqual(len(table_parts[0].rows), 1, "На первой странице должна быть одна строка таблицы")
        self.assertEqual(len(table_parts[1].rows), 1, "На второй странице — одна строка таблицы")
